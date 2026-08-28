"""Parse a résumé file into structured JSON.

This is extraction, not interpretation. Every string in the output appears
verbatim in the source document — nothing is summarized, normalized, or
inferred. That property is what later phases depend on: the fabrication guard
(§2.1) checks tailored output against this structure, so if the parser
paraphrased, the guard would be checking against a paraphrase.

Section detection is heuristic and deliberately conservative. Text that does
not clearly belong to a known section lands in `preamble` or the previous
section rather than being dropped — losing a line from someone's résumé is
worse than filing it imperfectly.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from pydantic import BaseModel, Field

#: Canonical section names and the headings that map to them. Order matters
#: only for reporting; matching is by longest heading first.
SECTION_PATTERNS: dict[str, tuple[str, ...]] = {
    "summary": ("summary", "objective", "profile", "about me", "professional summary"),
    "experience": (
        "experience",
        "work experience",
        "employment",
        "employment history",
        "professional experience",
        "work history",
    ),
    "education": ("education", "academic background", "academics"),
    "skills": ("skills", "technical skills", "core competencies", "technologies"),
    "projects": ("projects", "personal projects", "selected projects", "side projects"),
    "certifications": ("certifications", "certificates", "licenses"),
    "publications": ("publications", "papers"),
    "awards": ("awards", "honors", "achievements"),
    "languages": ("languages",),
    "interests": ("interests", "hobbies"),
}

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
# Two shapes, tried in order: an international number (leading +, country
# code, 7-14 more digits) and the North American 3-3-4 grouping. Résumés carry
# both, and a résumé from outside North America is not an edge case.
_PHONE_RE = re.compile(
    r"\+\d{1,3}[\s.\-()]{0,2}(?:\d[\s.\-()]{0,2}){6,13}\d"
    r"|\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}"
)
_URL_RE = re.compile(r"https?://[^\s,;)]+|(?:www\.|linkedin\.com|github\.com)[^\s,;)]+")

#: A heading line is short and has no sentence punctuation.
_MAX_HEADING_WORDS = 5


class ParseError(Exception):
    """The file could not be read as a résumé."""


class Contact(BaseModel):
    """Contact details found in the document, verbatim."""

    name: str | None = None
    email: str | None = None
    phone: str | None = None
    links: list[str] = Field(default_factory=list)


class ParsedResume(BaseModel):
    """A résumé as structure, with every string traceable to the source."""

    contact: Contact = Field(default_factory=Contact)
    #: Anything above the first recognized heading.
    preamble: list[str] = Field(default_factory=list)
    #: Canonical section name → its lines, in document order.
    sections: dict[str, list[str]] = Field(default_factory=dict)
    #: Every line, in order. The authority for "was this in the source?".
    raw_lines: list[str] = Field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(self.raw_lines)

    def section(self, name: str) -> list[str]:
        return self.sections.get(name, [])


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------


def extract_text(data: bytes, filename: str) -> str:
    """Pull plain text out of a résumé file.

    Supports PDF, DOCX, and plain text — the three things a résumé actually
    arrives as. A .doc (old binary Word) is rejected rather than mangled.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix in (".txt", ".md", ""):
        return data.decode("utf-8", errors="replace")
    if suffix == ".doc":
        raise ParseError("legacy .doc is not supported; save as .docx or PDF")
    raise ParseError(f"unsupported résumé format: {suffix or 'unknown'}")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 - any malformed PDF lands here
        raise ParseError(f"could not read PDF: {type(exc).__name__}") from exc

    text = "\n".join(pages)
    if not text.strip():
        raise ParseError(
            "no text found in PDF — it may be a scan. Export a text-based PDF; "
            "an ATS cannot read a scanned one either."
        )
    return text


#: Formatting properties whose change marks a run boundary as a field boundary.
#: Compared as serialized XML — the point is only whether two runs are styled
#: the same, not what the styling is.
def _run_style(run: object) -> str:
    props = getattr(run, "_r", None)
    if props is None:
        return ""
    found = props.find(f"{{{_W_NS}}}rPr")
    return "" if found is None else str(found.xml)


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _paragraph_text(paragraph: object) -> str:
    """A paragraph's text, with field boundaries preserved.

    Two separate defects, both observed on the owner's own résumé, both of
    which produce the same symptom — `California State University East
    BayHayward, CA`, `Master of Science in Business AnalyticsJan 2025 – Dec
    2026`. An ATS reads `AnalyticsJan` as one token, so the degree is wrong and
    the start date is not there at all.

    **A real tab is not a run.** It is a `<w:tab/>` element between runs, so
    `Paragraph.text` — which concatenates run text — drops it. Only tabs inside
    a run count: `<w:tabs>` under `<w:pPr>` is a *tab stop definition*, not a
    character, and treating it as one puts a space at the front of the
    paragraph and nothing where the separator belongs.

    **The rest of the time there is no tab at all.** The owner's education
    lines carry a right-aligned tab *stop* and two differently styled runs —
    `East Bay` in bold, `Hayward, CA` in italic — with nothing between them.
    Whatever produced the file lost the tab and kept the stop, so the fields
    are genuinely adjacent in the source.

    Reading is not enough to fix that, so this repairs it: a run boundary where
    the styling changes *and* the text goes from lowercase to uppercase is two
    fields, and gets a space. Both conditions are required. Either alone is
    common inside an ordinary word — a single emphasized letter, a `PostgreSQL`
    split across runs — and the repair would corrupt it.
    """
    from docx.text.paragraph import Paragraph

    assert isinstance(paragraph, Paragraph)

    out = ""
    previous_style: str | None = None
    for element in _runs_in_order(paragraph._p):
        text = "".join(
            " " if node.tag.rsplit("}", 1)[-1] in ("tab", "br", "cr") else (node.text or "")
            for node in element.iter()
            if node.tag.rsplit("}", 1)[-1] in ("t", "tab", "br", "cr")
        )
        if not text:
            continue
        found = element.find(f"{{{_W_NS}}}rPr")
        style = "" if found is None else str(found.xml)
        if (
            out
            and previous_style is not None
            and style != previous_style
            and out[-1].islower()
            and text[0].isupper()
        ):
            out += " "
        out += text
        previous_style = style
    return out.strip()


def _runs_in_order(element: object) -> list:  # type: ignore[type-arg]
    """Every run in a paragraph, including the ones inside hyperlinks.

    `Paragraph.runs` returns only direct `<w:r>` children. A link is a
    `<w:hyperlink>` wrapping its own runs, so reading `runs` alone silently
    drops it — which cost the owner's résumé both its LinkedIn and its GitHub
    URL, the two links a recruiter actually clicks.
    """
    runs = []
    for child in element:  # type: ignore[attr-defined]
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "r":
            runs.append(child)
        elif tag == "hyperlink":
            runs.extend(c for c in child if c.tag.rsplit("}", 1)[-1] == "r")
    return runs


def _extract_docx(data: bytes) -> str:
    """Text from a DOCX, in the order the document actually reads.

    Body children are walked in document order rather than taking
    `document.paragraphs` and then `document.tables`. Those two properties each
    return their own kind in order, so concatenating them moves every table to
    the end of the file regardless of where it sits on the page.

    That is not cosmetic — section assignment is positional. On the owner's
    résumé the Skills block is a two-column table under a `TECHNICAL SKILLS`
    heading near the top; appended last, it landed under whatever heading came
    last instead, which was `PROJECTS`. The result: a Skills section holding
    only certifications, six skills lines filed as project bullets, and the
    tailorer then rewriting `Frameworks & Web  FastAPI, React, ...` as if it
    were a sentence about a project.
    """
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ParseError(f"could not read DOCX: {type(exc).__name__}") from exc

    lines: list[str] = []
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            lines.append(_paragraph_text(Paragraph(child, document)))
        elif tag == "tbl":
            for row in Table(child, document).rows:
                # A résumé table is a layout device — a label column and a
                # value column. Joined with two spaces so the label stays
                # attached to what it labels, which is what makes
                # `Frameworks & Web  FastAPI, React` one readable line.
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append("  ".join(dict.fromkeys(cells)))
    return "\n".join(lines)


# --------------------------------------------------------------------------
# Structure
# --------------------------------------------------------------------------


def _normalize_heading(line: str) -> str:
    return re.sub(r"[^a-z ]", "", line.lower()).strip()


def _match_section(line: str) -> str | None:
    """Return the canonical section this line is a heading for, if any."""
    stripped = line.strip()
    if not stripped or len(stripped.split()) > _MAX_HEADING_WORDS:
        return None
    # A line ending in a period is prose, not a heading.
    if stripped.endswith((".", ",", ";")):
        return None

    normalized = _normalize_heading(stripped)
    if not normalized:
        return None

    for canonical, headings in SECTION_PATTERNS.items():
        if normalized in headings:
            return canonical

    return _match_compound_heading(stripped, normalized)


#: Punctuation that appears in résumé *content* and never in a section
#: heading. A heading names a section: it carries no list, no annotation, no
#: date and no title/subtitle break.
#:
#: Load-bearing. Without it a skills row like `Languages  Python, Rust` is
#: Title Case, short, and contains the word `languages` — so it matched the
#: Languages section and took the rest of the résumé's skills with it. The
#: owner's own row escaped only by being too long for the word limit.
_NOT_IN_A_HEADING = ",[]()0123456789—–•|/:"


def _looks_like_a_heading(line: str) -> bool:
    """All-caps, or Title Case With Every Word Capitalized.

    The two ways a résumé writes a heading, and the gate on the looser match
    below. A line of prose fails both.
    """
    if any(c in _NOT_IN_A_HEADING for c in line):
        return False
    letters = [c for c in line if c.isalpha()]
    if not letters:
        return False
    if all(c.isupper() for c in letters):
        return True
    return all(word[0].isupper() for word in line.split() if word and word[0].isalpha())


def _match_compound_heading(stripped: str, normalized: str) -> str | None:
    """A heading that names its section and then says more.

    `CERTIFICATIONS & TRAINING` normalizes to `certifications training`, which
    is not any entry in `SECTION_PATTERNS`, so exact matching filed every
    certification below it under whichever section preceded it. On the owner's
    résumé that was Skills, which then held no skills at all.
    `ACHIEVEMENTS & ACTIVITIES` did the same thing to Projects, and the
    tailorer went on to rewrite `School-Level Championships in Chess, Football,
    and Cricket` as though it were a project bullet.

    Extending the pattern lists instead would be endless — `AWARDS & HONORS`,
    `PROJECTS AND RESEARCH`, `EDUCATION & TRAINING` — and each omission fails
    the same silent way.

    Two conditions, and both are needed. The line must *look* like a heading,
    which is what keeps this off ordinary prose; and a known heading must
    appear in it as whole words, so `Improved certification coverage` does not
    match on a shared stem. The most specific match wins, so `WORK EXPERIENCE`
    is not decided by whichever of `experience` and `work experience` is
    reached first.
    """
    if not _looks_like_a_heading(stripped):
        return None

    words = set(normalized.split())
    best: tuple[int, str] | None = None
    for canonical, headings in SECTION_PATTERNS.items():
        for heading in headings:
            parts = heading.split()
            if not set(parts) <= words:
                continue
            if best is None or len(parts) > best[0]:
                best = (len(parts), canonical)
    return best[1] if best else None


def _find_contact(lines: list[str]) -> Contact:
    """Contact details, taken verbatim from wherever they appear.

    The name heuristic is the first non-empty line that is not itself contact
    data — conventional for résumés, and wrong often enough that it stays
    editable in the profile rather than being treated as authoritative.
    """
    blob = "\n".join(lines)

    email_match = _EMAIL_RE.search(blob)
    phone_match = _PHONE_RE.search(blob)

    links: list[str] = []
    for match in _URL_RE.finditer(blob):
        url = match.group(0).rstrip(".,;")
        if url not in links:
            links.append(url)

    name = None
    for line in lines[:5]:
        candidate = line.strip()
        if not candidate:
            continue
        if _EMAIL_RE.search(candidate) or _PHONE_RE.search(candidate):
            continue
        if _URL_RE.search(candidate):
            continue
        if len(candidate.split()) <= 5:
            name = candidate
            break

    return Contact(
        name=name,
        email=email_match.group(0) if email_match else None,
        phone=phone_match.group(0).strip() if phone_match else None,
        links=links,
    )


def parse_text(text: str) -> ParsedResume:
    """Segment already-extracted text into sections."""
    raw_lines = [line.rstrip() for line in text.splitlines()]
    non_empty = [line for line in raw_lines if line.strip()]

    resume = ParsedResume(contact=_find_contact(non_empty), raw_lines=non_empty)

    current: str | None = None
    for line in non_empty:
        matched = _match_section(line)
        if matched is not None:
            current = matched
            resume.sections.setdefault(current, [])
            continue

        if current is None:
            resume.preamble.append(line)
        else:
            resume.sections[current].append(line)

    return resume


def parse_resume(data: bytes, filename: str) -> ParsedResume:
    """Extract and segment a résumé file."""
    return parse_text(extract_text(data, filename))
