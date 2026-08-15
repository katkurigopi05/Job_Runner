"""Résumé parsing and assembly.

The parser's contract is extraction, not interpretation: every string it
returns must appear in the source. Phase 3's fabrication guard checks tailored
output against this structure, so a parser that paraphrased would quietly
undermine the guard.
"""

from __future__ import annotations

import io

import pytest
from pypdf import PdfReader

from packages.tailor.assemble import AssemblyOptions, assemble_html, assemble_pdf, describe
from packages.tailor.parse import ParsedResume, ParseError, extract_text, parse_text
from packages.tailor.projects import LinkStyle
from tests.test_projects import make_project

SAMPLE = """Ada Lovelace
ada@example.com | +1 (555) 555-0100 | github.com/ada | linkedin.com/in/ada

Summary
Backend engineer focused on distributed systems and developer tooling.

Experience
Staff Engineer, Analytical Engines Ltd — 2021 to present
Designed the note-taking subsystem handling 2M events per day.
Mentored four engineers through promotion.

Education
BSc Mathematics, University of London

Skills
Python, PostgreSQL, Playwright, Docker
"""


def test_extracts_contact_details() -> None:
    resume = parse_text(SAMPLE)

    assert resume.contact.name == "Ada Lovelace"
    assert resume.contact.email == "ada@example.com"
    assert resume.contact.phone is not None
    assert any("github.com/ada" in link for link in resume.contact.links)


def test_segments_known_sections() -> None:
    resume = parse_text(SAMPLE)

    assert set(resume.sections) >= {"summary", "experience", "education", "skills"}
    assert "Backend engineer" in resume.section("summary")[0]
    assert len(resume.section("experience")) == 3


def test_every_line_is_traceable_to_the_source() -> None:
    """The property Phase 3's fabrication guard will depend on."""
    resume = parse_text(SAMPLE)
    source_lines = {line.strip() for line in SAMPLE.splitlines() if line.strip()}

    for line in resume.raw_lines:
        assert line.strip() in source_lines

    for lines in resume.sections.values():
        for line in lines:
            assert line.strip() in source_lines


def test_unknown_sections_are_kept_not_dropped() -> None:
    text = SAMPLE + "\nVolunteering\nTaught Python at a local library.\n"
    resume = parse_text(text)
    assert "Taught Python at a local library." in resume.raw_lines


def test_headings_are_matched_case_insensitively() -> None:
    resume = parse_text("WORK EXPERIENCE\nDid a thing.\n")
    assert resume.section("experience") == ["Did a thing."]


def test_prose_is_not_mistaken_for_a_heading() -> None:
    """A short sentence ending in a period is content, not a section."""
    resume = parse_text("Summary\nI like skills.\nMore text.\n")
    assert "I like skills." in resume.section("summary")


def test_content_before_any_heading_lands_in_preamble() -> None:
    resume = parse_text("Ada Lovelace\nada@example.com\n\nSummary\nHello.\n")
    assert "Ada Lovelace" in resume.preamble


def test_empty_document_parses_to_empty_structure() -> None:
    resume = parse_text("")
    assert resume.raw_lines == []
    assert resume.sections == {}


# --------------------------------------------------------------------------
# File formats
# --------------------------------------------------------------------------


def test_extract_from_txt() -> None:
    assert "Ada Lovelace" in extract_text(SAMPLE.encode(), "resume.txt")


def test_extract_from_pdf() -> None:
    from weasyprint import HTML

    pdf = HTML(string="<p>Ada Lovelace</p><p>ada@example.com</p>").write_pdf()
    text = extract_text(bytes(pdf), "resume.pdf")
    assert "Ada Lovelace" in text


def test_extract_from_docx() -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("Experience")
    buf = io.BytesIO()
    document.save(buf)

    text = extract_text(buf.getvalue(), "resume.docx")
    assert "Ada Lovelace" in text
    assert "Experience" in text


def test_docx_tables_are_not_lost() -> None:
    """Résumé templates lean on tables; their content still has to survive."""
    import docx

    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "PostgreSQL"
    buf = io.BytesIO()
    document.save(buf)

    text = extract_text(buf.getvalue(), "resume.docx")
    assert "Python" in text
    assert "PostgreSQL" in text


def test_legacy_doc_is_rejected_clearly() -> None:
    with pytest.raises(ParseError, match="save as .docx"):
        extract_text(b"\xd0\xcf", "resume.doc")


def test_unsupported_format_is_rejected() -> None:
    with pytest.raises(ParseError, match="unsupported"):
        extract_text(b"x", "resume.pages")


def test_scanned_pdf_is_reported_not_silently_empty() -> None:
    """An image-only PDF is unreadable to an ATS too; say so."""
    from weasyprint import HTML

    pdf = HTML(string="<div style='height:2in'></div>").write_pdf()
    with pytest.raises(ParseError, match="scan"):
        extract_text(bytes(pdf), "scan.pdf")


def test_corrupt_pdf_raises_parse_error() -> None:
    with pytest.raises(ParseError, match="could not read PDF"):
        extract_text(b"not a pdf at all", "resume.pdf")


# --------------------------------------------------------------------------
# Assembly — projects rebuilt into every résumé
# --------------------------------------------------------------------------


def _extract(pdf: bytes) -> tuple[str, list[str]]:
    reader = PdfReader(io.BytesIO(pdf))
    page = reader.pages[0]
    urls = []
    for annot in page.get("/Annots") or []:
        action = annot.get_object().get("/A")
        if action and action.get("/URI"):
            urls.append(str(action["/URI"]))
    return page.extract_text() or "", urls


def test_assembly_keeps_source_text_verbatim() -> None:
    resume = parse_text(SAMPLE)
    html = assemble_html(resume, [])
    assert "Designed the note-taking subsystem handling 2M events per day." in html


def test_projects_are_added_to_the_resume() -> None:
    resume = parse_text(SAMPLE)
    projects = [make_project(name="jobrunner", description="Local job-application agent")]

    html = assemble_html(resume, projects)

    assert "jobrunner" in html
    assert "Local job-application agent" in html
    assert 'href="https://github.com/octocat/jobrunner"' in html


def test_generated_projects_replace_a_source_projects_section() -> None:
    """Otherwise the résumé prints two Projects sections."""
    text = SAMPLE + "\nProjects\nAn old project I listed by hand.\n"
    resume = parse_text(text)
    html = assemble_html(resume, [make_project(name="fresh")])

    assert html.count("<h2>Projects</h2>") == 1
    assert "An old project I listed by hand." not in html


def test_source_projects_survive_when_none_are_supplied() -> None:
    text = SAMPLE + "\nProjects\nAn old project I listed by hand.\n"
    resume = parse_text(text)
    html = assemble_html(resume, [])
    assert "An old project I listed by hand." in html


def test_assembled_pdf_carries_project_links() -> None:
    resume = parse_text(SAMPLE)
    projects = [make_project(name="jobrunner", url="https://github.com/octocat/jobrunner")]

    text, urls = _extract(assemble_pdf(resume, projects))

    assert "Ada Lovelace" in text
    assert "github.com/octocat/jobrunner" in text  # ATS-visible
    assert "https://github.com/octocat/jobrunner" in urls  # clickable


def test_icon_only_assembly_still_links_for_humans() -> None:
    """Icon-only drops the project URL from the text layer, but keeps it
    clickable. The contact line's own links are unaffected."""
    resume = parse_text(SAMPLE)
    options = AssemblyOptions(link_style=LinkStyle.ICON_ONLY)
    text, urls = _extract(assemble_pdf(resume, [make_project()], options))

    assert "github.com/octocat/jobrunner" not in text
    assert "https://github.com/octocat/jobrunner" in urls


def test_projects_can_be_turned_off() -> None:
    resume = parse_text(SAMPLE)
    options = AssemblyOptions(include_projects=False)
    html = assemble_html(resume, [make_project(name="jobrunner")], options)
    assert "jobrunner" not in html


def test_describe_summarizes_without_rendering() -> None:
    resume = parse_text(SAMPLE)
    report = describe(resume, [make_project(name="jobrunner")])

    assert "experience" in report.sections
    assert report.project_names == ["jobrunner"]
    assert report.source_line_count > 0


def test_assembly_escapes_source_text() -> None:
    resume = ParsedResume(sections={"summary": ["<script>alert(1)</script>"]})
    html = assemble_html(resume, [])
    assert "<script>" not in html


@pytest.mark.parametrize(
    "number",
    [
        "+1 (555) 555-0100",
        "555-555-0100",
        "+44 20 7946 0958",
        "+91 98765 43210",
        "+91-98765-43210",
    ],
)
def test_phone_formats_are_recognized(number: str) -> None:
    """A résumé from outside North America is not an edge case."""
    resume = parse_text(f"Ada Lovelace\nada@example.com | {number}\n")
    assert resume.contact.phone is not None
