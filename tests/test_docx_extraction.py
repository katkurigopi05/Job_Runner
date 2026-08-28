"""DOCX extraction, against the three defects found on the owner's own résumé.

All three were silent. The parse succeeded, every section came back non-empty,
and the tailorer ran to completion — on content filed under the wrong heading.
Each test here builds a .docx with the same shape and asserts the text that
comes out, because none of these are visible from the parsed structure alone.
"""

from __future__ import annotations

import io

import docx
import pytest

from packages.tailor.parse import extract_text, parse_text


def _render(document: docx.document.Document) -> str:
    buffer = io.BytesIO()
    document.save(buffer)
    return extract_text(buffer.getvalue(), "resume.docx")


# --------------------------------------------------------------------------
# Document order
# --------------------------------------------------------------------------


def test_a_table_is_read_where_it_sits_not_at_the_end() -> None:
    """The defect that put six skills lines inside the Projects section.

    `document.paragraphs` and `document.tables` each return their own kind in
    order, so reading one then the other moves every table to the end of the
    file. The owner's Skills block is a two-column table near the top; appended
    last, it inherited the last heading in the document instead — `PROJECTS`.
    """
    document = docx.Document()
    document.add_paragraph("TECHNICAL SKILLS")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Languages"
    table.rows[0].cells[1].text = "Python, Rust"
    document.add_paragraph("PROJECTS")
    document.add_paragraph("Built a thing.")

    lines = [line for line in _render(document).splitlines() if line.strip()]

    assert lines.index("TECHNICAL SKILLS") < lines.index("Languages  Python, Rust")
    assert lines.index("Languages  Python, Rust") < lines.index("PROJECTS")


def test_a_table_row_lands_in_the_section_it_appears_under() -> None:
    document = docx.Document()
    document.add_paragraph("TECHNICAL SKILLS")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Languages"
    table.rows[0].cells[1].text = "Python, Rust"
    document.add_paragraph("PROJECTS")
    document.add_paragraph("Built a thing.")

    parsed = parse_text(_render(document))

    assert any("Python, Rust" in line for line in parsed.section("skills"))
    assert not any("Python, Rust" in line for line in parsed.section("projects"))


# --------------------------------------------------------------------------
# Field separation
# --------------------------------------------------------------------------


def test_a_real_tab_between_runs_becomes_a_space() -> None:
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Software Engineer")
    paragraph.add_run().add_tab()
    paragraph.add_run("Jan 2021 - Present")

    assert "Software Engineer Jan 2021 - Present" in _render(document)


def test_a_tab_stop_definition_is_not_read_as_a_character() -> None:
    """`<w:tabs>` under `<w:pPr>` positions text; it is not text.

    Counting it put a space at the front of the paragraph and none where the
    separator actually belonged.
    """
    document = docx.Document()
    paragraph = document.add_paragraph("EDUCATION")
    paragraph.paragraph_format.tab_stops.add_tab_stop(docx.shared.Inches(7))

    assert _render(document).splitlines()[0] == "EDUCATION"


def test_differently_styled_runs_are_separated_at_a_case_change() -> None:
    """The owner's education lines: two runs, different styling, no tab at all.

    Whatever produced the file kept the right-aligned tab *stop* and lost the
    tab, so `East Bay` and `Hayward, CA` are genuinely adjacent in the source.
    An ATS reads `BayHayward` as one token.
    """
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("California State University East Bay").bold = True
    paragraph.add_run("Hayward, CA").italic = True

    assert "East Bay Hayward, CA" in _render(document)


def test_a_degree_is_separated_from_its_date() -> None:
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Master of Science in Business Analytics").bold = True
    paragraph.add_run("Jan 2025 - Dec 2026").italic = True

    assert "Business Analytics Jan 2025" in _render(document)


def test_runs_styled_the_same_are_never_separated() -> None:
    """Both conditions are required, or an ordinary word gets split.

    A run boundary alone is not a field boundary — Word splits runs for
    spell-check state and revision marks — so `PostgreSQL` broken across two
    identically styled runs must survive intact.
    """
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Postgre")
    paragraph.add_run("SQL")

    assert "PostgreSQL" in _render(document)


def test_a_case_change_alone_does_not_separate() -> None:
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hugging")
    paragraph.add_run("Face")

    assert "HuggingFace" in _render(document)


# --------------------------------------------------------------------------
# Hyperlinks
# --------------------------------------------------------------------------


def test_a_hyperlink_is_not_dropped() -> None:
    """`Paragraph.runs` returns only direct children, so a link is invisible.

    Reading runs alone cost the owner's résumé both its LinkedIn and its GitHub
    URL — the two links a recruiter clicks.
    """
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("me@example.com | ")
    paragraph.add_hyperlink = None  # not part of the API; build the element by hand
    _add_hyperlink(paragraph, "github.com/janedoe")

    text = _render(document)
    assert "github.com/janedoe" in text

    parsed = parse_text(text)
    assert "github.com/janedoe" in parsed.contact.links


def _add_hyperlink(paragraph: docx.text.paragraph.Paragraph, text: str) -> None:
    """Append a `<w:hyperlink>` wrapping one run, the way Word writes a link."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    link = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)
    assert link.tag == qn("w:hyperlink")


# --------------------------------------------------------------------------
# Unsupported input
# --------------------------------------------------------------------------


def test_a_corrupt_docx_is_reported_not_mangled() -> None:
    from packages.tailor.parse import ParseError

    with pytest.raises(ParseError):
        extract_text(b"not a docx at all", "resume.docx")
